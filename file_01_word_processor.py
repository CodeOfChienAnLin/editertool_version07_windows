"""
Word文件處理相關功能模組
"""
import os
import io
import traceback
import platform
import tkinter as tk
from tkinter import messagebox, simpledialog
import docx2txt
import msoffcrypto
from docx import Document
from io import BytesIO
import threading

# --- COM 相關匯入和檢查 (用於 Windows + Word 解析) ---
HAS_PYWIN32 = False
if platform.system() == 'Windows': # COM 只在 Windows 上可用
    try:
        import win32com.client as win32
        import pythoncom # 需要初始化 COM
        HAS_PYWIN32 = True
    except ImportError:
        print("警告：未找到 pywin32 模組。COM 功能需要 Windows、Microsoft Word 和 'pip install pywin32'。")
        pass # 即使沒有 pywin32，程式仍可嘗試啟動 (但功能受限)

def load_and_display_word_content(self, file_path, password=None, skip_processing_window=False):
    """
    載入並顯示 Word 文件內容，整合 COM 和原有方法，並提取圖片。

    參數:
        file_path: Word檔案路徑
        password: 檔案密碼（如果有的話）
        skip_processing_window: 是否跳過顯示處理中提示視窗
    
    """
    try:
        # 更新狀態欄，提示使用者正在開啟檔案
        self.status_bar.config(text=f"正在開啟 Word 檔案: {os.path.basename(file_path)}...")
        
        # 創建處理中提示視窗（除非指定跳過）
        processing_window = None
        if not skip_processing_window:
            processing_window = tk.Toplevel(self.root)
            processing_window.title("處理中")
            processing_window.geometry("300x100")
            processing_window.resizable(False, False)
            processing_window.transient(self.root)
            
            # 居中顯示
            processing_window.update_idletasks()
            width = processing_window.winfo_width()
            height = processing_window.winfo_height()
            x = (processing_window.winfo_screenwidth() // 2) - (width // 2)
            y = (processing_window.winfo_screenheight() // 2) - (height // 2)
            processing_window.geometry(f"+{x}+{y}")
            
            # 提示訊息
            tk.Label(processing_window, text=f"正在開啟 Word 檔案...\n請稍候", pady=20).pack()
            
            # 強制更新視窗
            processing_window.update()
        
        # 清空圖片區域
        from file_02_image_handler import clear_images
        clear_images(self)
        
        # 嘗試使用 COM 解析 (僅在 Windows 上)
        if HAS_PYWIN32 and platform.system() == 'Windows':
            try:
                content = parse_word_document_com(self, file_path)
                if content:
                    # 更新文字區域
                    self.text_area.delete("1.0", tk.END)
                    
                    # 先進行文字修正和簡體字檢查
                    from text_01_correction import correct_text_for_word_import
                    corrected_content = correct_text_for_word_import(self, content)
                    
                    # 插入修正後的文字
                    self.text_area.insert("1.0", corrected_content)
                    
                    # 提取圖片
                    from file_02_image_handler import extract_images_from_docx
                    extract_images_from_docx(self, file_path)
                    
                    # 關閉處理中提示視窗
                    if processing_window:
                        processing_window.destroy()
                    
                    # 更新狀態欄
                    self.status_bar.config(text=f"已載入檔案: {file_path}")
                    return True
            except Exception as com_error:
                print(f"COM 解析錯誤: {com_error}")
        
        # 如果 COM 解析失敗或不可用，使用內部處理方法
        content = process_word_file_internal(self, file_path, password)
        if content:
            # 更新文字區域
            self.text_area.delete("1.0", tk.END)
            
            # 先進行文字修正和簡體字檢查
            from text_01_correction import correct_text_for_word_import
            corrected_content = correct_text_for_word_import(self, content)
            
            # 插入修正後的文字
            self.text_area.insert("1.0", corrected_content)
            
            # 提取圖片
            from file_02_image_handler import extract_images_from_docx
            extract_images_from_docx(self, file_path)
            
            # 關閉處理中提示視窗
            if processing_window:
                processing_window.destroy()
            
            # 更新狀態欄
            self.status_bar.config(text=f"已載入檔案: {file_path}")
            return True
        
        # 如果所有方法都失敗
        if processing_window:
            processing_window.destroy()
        messagebox.showerror("錯誤", f"無法讀取檔案 '{os.path.basename(file_path)}'。\n兩種解析方法均失敗。")
        self.status_bar.config(text=f"讀取檔案失敗: {os.path.basename(file_path)}")
        return False
        
    except Exception as e:
        # 關閉處理中提示視窗
        if 'processing_window' in locals() and processing_window:
            processing_window.destroy()
            
        error_msg = f"處理 Word 文件時發生錯誤: {str(e)}"
        print(error_msg)
        messagebox.showerror("錯誤", error_msg)
        self.status_bar.config(text="處理檔案時發生錯誤")
        return False

def process_word_file_internal(self, file_path, password=None):
    """
    處理Word檔案（內部方法，優先使用 docx2txt，回退到 python-docx）。
    如果提供密碼，則先解密。
    **修改：** 在解密成功後調用圖片提取。

    參數:
        file_path: Word檔案路徑
        password: 檔案密碼（如果有的話）

    回傳:
        檔案內容
    """
    try:
        # 嘗試直接處理檔案（假設未加密）
        return _process_unencrypted_file(self, file_path)
    
    except Exception as e:
        error_message = str(e)
        
        # 檢查是否為密碼保護錯誤
        if _is_password_error(self, error_message):
            # 如果已提供密碼但仍然失敗
            if password:
                messagebox.showerror("錯誤", f"密碼不正確或檔案已損壞: {error_message}")
                return None
            
            # 否則，嘗試處理密碼保護的檔案
            return handle_password_protected_file(self, file_path)
        
        # 其他錯誤
        messagebox.showerror("錯誤", f"處理檔案時發生錯誤: {error_message}")
        return None

def _process_unencrypted_file(self, file_path):
    """處理未加密的Word檔案

    參數:
        file_path: Word檔案路徑

    回傳:
        檔案內容
    """
    try:
        # 首先嘗試使用 docx2txt
        content = docx2txt.process(file_path)
        return content
    
    except Exception as docx2txt_error:
        print(f"docx2txt 處理失敗: {str(docx2txt_error)}")
        
        # 如果 docx2txt 失敗，嘗試使用 python-docx
        try:
            doc = Document(file_path)
            return _extract_text_from_document(self, doc)
        
        except Exception as docx_error:
            # 兩種方法都失敗，拋出異常
            raise Exception(f"無法讀取檔案: {str(docx_error)}")

def _is_password_error(self, error_message):
    """檢查錯誤訊息是否與密碼保護相關

    參數:
        error_message: 錯誤訊息

    回傳:
        是否為密碼相關錯誤
    """
    password_error_keywords = [
        "password", "密碼", "protected", "保護", "encrypted", "加密",
        "Encryption", "encryptionInfo", "EncryptedPackage"
    ]
    
    return any(keyword in error_message for keyword in password_error_keywords)

def _extract_text_from_document(self, doc):
    """從 python-docx Document 物件中提取文字

    參數:
        doc: python-docx Document 物件

    回傳:
        提取的文字
    """
    full_text = []
    
    # 提取段落文字
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # 提取表格文字
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text)
            full_text.append(" | ".join(row_text))
    
    return "\n".join(full_text)

def handle_password_protected_file(self, file_path):
    """處理有密碼保護的Word檔案

    參數:
        file_path: 加密Word檔案的路徑
    
    回傳:
        解密後的檔案內容，如果失敗則返回 None
    """
    # 顯示密碼輸入對話框
    password = ask_password(self)
    
    if not password:
        self.status_bar.config(text="已取消密碼輸入")
        return None  # 使用者取消輸入
    
    # 更新狀態欄
    self.status_bar.config(text=f"正在處理加密檔案: {os.path.basename(file_path)}...")
    
    # 創建處理中提示視窗
    processing_window = tk.Toplevel(self.root)
    processing_window.title("處理中")
    processing_window.geometry("300x100")
    processing_window.resizable(False, False)
    processing_window.transient(self.root)
    
    # 居中顯示
    processing_window.update_idletasks()
    width = processing_window.winfo_width()
    height = processing_window.winfo_height()
    x = (processing_window.winfo_screenwidth() // 2) - (width // 2)
    y = (processing_window.winfo_screenheight() // 2) - (height // 2)
    processing_window.geometry(f"+{x}+{y}")
    
    # 提示訊息
    tk.Label(processing_window, text=f"正在處理加密檔案...\n請稍候", pady=20).pack()
    
    # 強制更新視窗
    processing_window.update()
    
    # 清空圖片區域
    from file_02_image_handler import clear_images
    clear_images(self)
    
    # 在背景執行解密和處理
    def process_encrypted_file_thread():
        try:
            # 創建一個臨時文件來存儲解密後的內容
            temp_file = BytesIO()
            
            # 打開加密檔案
            with open(file_path, 'rb') as file:
                office_file = msoffcrypto.OfficeFile(file)
                
                # 設置密碼
                office_file.load_key(password=password)
                
                # 解密到臨時文件
                office_file.decrypt(temp_file)
            
            # 創建臨時實體檔案以便 COM 處理
            import tempfile
            temp_docx_path = None
            
            try:
                # 重置臨時文件指針
                temp_file.seek(0)
                
                # 創建臨時實體檔案
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
                    temp_docx.write(temp_file.read())
                    temp_docx_path = temp_docx.name
                
                # 在主線程中使用 COM 解析
                self.root.after(0, lambda: process_with_com(temp_docx_path))
                
            except Exception as e:
                # 如果創建臨時文件失敗，使用內部方法解析
                error_message = str(e)
                print(f"創建臨時檔案失敗: {error_message}")
                
                # 重置臨時文件指針
                temp_file.seek(0)
                
                # 使用內部方法處理
                content = None
                try:
                    # 嘗試使用 docx2txt 處理解密後的內容
                    content = docx2txt.process(temp_file)
                except Exception as docx2txt_error:
                    print(f"docx2txt 處理解密檔案失敗: {str(docx2txt_error)}")
                    
                    # 如果 docx2txt 失敗，嘗試使用 python-docx
                    try:
                        temp_file.seek(0)  # 重置文件指針
                        doc = Document(temp_file)
                        content = _extract_text_from_document(self, doc)
                    except Exception as docx_error:
                        # 兩種方法都失敗，拋出異常
                        raise Exception(f"無法讀取解密後的檔案: {str(docx_error)}")
                
                # 在主線程中更新 UI
                self.root.after(0, lambda: update_ui_with_content(content))
            
        except Exception as e:
            # 如果解密失敗，可能是密碼錯誤
            error_message = str(e)
            self.root.after(0, lambda: handle_error(error_message))
    
    # 使用 COM 處理解密後的檔案
    def process_with_com(temp_docx_path):
        try:
            # 使用 COM 解析 (僅在 Windows 上)
            if HAS_PYWIN32 and platform.system() == 'Windows':
                try:
                    # 使用 COM 解析臨時檔案
                    content = parse_word_document_com(self, temp_docx_path)
                    if content:
                        # 更新 UI
                        update_ui_with_content(content)
                        
                        # 提取圖片
                        try:
                            from file_02_image_handler import extract_images_from_docx
                            extract_images_from_docx(self, temp_docx_path)
                        except Exception as img_error:
                            print(f"提取圖片時出錯: {str(img_error)}")
                        
                        # 刪除臨時檔案
                        try:
                            os.remove(temp_docx_path)
                        except:
                            pass
                        
                        return
                except Exception as com_error:
                    print(f"COM 解析錯誤: {com_error}")
            
            # 如果 COM 解析失敗或不可用，使用 load_and_display_word_content 處理
            try:
                # 使用 load_and_display_word_content 處理，但跳過顯示處理中提示視窗
                load_and_display_word_content(self, temp_docx_path, password=None, skip_processing_window=True)
                
                # 刪除臨時檔案
                try:
                    os.remove(temp_docx_path)
                except:
                    pass
                
                # 更新狀態欄
                self.status_bar.config(text=f"已載入加密檔案: {os.path.basename(file_path)}")
                
                # 關閉處理中提示視窗
                processing_window.destroy()
                
                return
            except Exception as e:
                print(f"使用 load_and_display_word_content 處理失敗: {str(e)}")
                
                # 如果失敗，繼續使用內部方法處理
                with open(temp_docx_path, 'rb') as file:
                    temp_file = BytesIO(file.read())
                
                content = None
                try:
                    # 嘗試使用 docx2txt 處理解密後的內容
                    content = docx2txt.process(temp_file)
                except Exception as docx2txt_error:
                    print(f"docx2txt 處理解密檔案失敗: {str(docx2txt_error)}")
                    
                    # 如果 docx2txt 失敗，嘗試使用 python-docx
                    try:
                        temp_file.seek(0)  # 重置文件指針
                        doc = Document(temp_file)
                        content = _extract_text_from_document(self, doc)
                    except Exception as docx_error:
                        # 兩種方法都失敗，拋出異常
                        raise Exception(f"無法讀取解密後的檔案: {str(docx_error)}")
                
                # 更新 UI
                update_ui_with_content(content)
                
                # 提取圖片
                try:
                    from file_02_image_handler import extract_images_from_docx
                    extract_images_from_docx(self, temp_docx_path)
                except Exception as img_error:
                    print(f"提取圖片時出錯: {str(img_error)}")
                
                # 刪除臨時檔案
                try:
                    os.remove(temp_docx_path)
                except:
                    pass
            
        except Exception as e:
            # 處理錯誤
            processing_window.destroy()
            self.status_bar.config(text=f"處理加密檔案時出錯")
            messagebox.showerror("錯誤", f"處理加密檔案時發生錯誤: {str(e)}")
    
    # 在主線程中更新 UI
    def update_ui_with_content(content):
        try:
            # 關閉處理中提示視窗
            processing_window.destroy()
            
            if content:
                # 更新文字區域
                self.text_area.delete("1.0", tk.END)
                
                # 先進行文字修正和簡體字檢查
                from text_01_correction import correct_text_for_word_import
                corrected_content = correct_text_for_word_import(self, content)
                
                # 插入修正後的文字
                self.text_area.insert("1.0", corrected_content)
                
                # 更新狀態欄
                self.status_bar.config(text=f"已載入加密檔案: {os.path.basename(file_path)}")
            else:
                # 更新狀態欄
                self.status_bar.config(text="無法讀取加密檔案內容")
                messagebox.showerror("錯誤", "無法讀取加密檔案內容")
        except Exception as ui_error:
            # 更新狀態欄
            self.status_bar.config(text=f"處理加密檔案時出錯: {str(ui_error)}")
            messagebox.showerror("錯誤", f"處理加密檔案時出錯: {str(ui_error)}")
    
    # 處理錯誤
    def handle_error(error_message):
        # 關閉處理中提示視窗
        processing_window.destroy()
        
        # 如果解密失敗，可能是密碼錯誤
        if "Failed to verify password" in error_message or "密碼" in error_message:
            # 顯示錯誤訊息
            messagebox.showwarning("警告", "密碼不正確，請重新輸入")
            
            # 遞歸調用，再次嘗試
            handle_password_protected_file(self, file_path)
        else:
            # 其他錯誤
            self.status_bar.config(text=f"處理加密檔案時出錯")
            messagebox.showerror("錯誤", f"處理加密檔案時發生錯誤: {error_message}")
    
    # 啟動背景執行緒
    threading.Thread(target=process_encrypted_file_thread).start()
    
    return None  # 由背景執行緒處理後續操作

def ask_password(self):
    """顯示密碼輸入對話框

    回傳:
        使用者輸入的密碼
    """
    # 創建一個自定義對話框
    dialog = tk.Toplevel(self.root)
    dialog.title("輸入密碼")
    dialog.geometry("300x150")
    dialog.resizable(False, False)
    dialog.transient(self.root)  # 設置為主窗口的子窗口
    dialog.grab_set()  # 模態對話框
    
    # 居中顯示
    dialog.update_idletasks()
    width = dialog.winfo_width()
    height = dialog.winfo_height()
    x = (dialog.winfo_screenwidth() // 2) - (width // 2)
    y = (dialog.winfo_screenheight() // 2) - (height // 2)
    dialog.geometry(f"+{x}+{y}")
    
    # 說明標籤
    tk.Label(dialog, text="此檔案已加密，請輸入密碼：", pady=10).pack()
    
    # 密碼輸入框 - 移除 show="*" 參數，使密碼以明碼顯示
    password_var = tk.StringVar()
    password_entry = tk.Entry(dialog, textvariable=password_var, width=30)
    password_entry.pack(pady=5)
    password_entry.focus_set()  # 設置焦點
    
    # 結果變數
    result = [None]  # 使用列表以便在內部函數中修改
    
    # 確定按鈕事件
    def on_ok():
        result[0] = password_var.get()
        dialog.destroy()
    
    # 取消按鈕事件
    def on_cancel():
        dialog.destroy()
    
    # 按鈕框架
    button_frame = tk.Frame(dialog)
    button_frame.pack(pady=10)
    
    # 確定按鈕
    ok_button = tk.Button(button_frame, text="確定", width=10, command=on_ok)
    ok_button.pack(side=tk.LEFT, padx=5)
    
    # 取消按鈕
    cancel_button = tk.Button(button_frame, text="取消", width=10, command=on_cancel)
    cancel_button.pack(side=tk.LEFT, padx=5)
    
    # 綁定回車鍵
    dialog.bind("<Return>", lambda event: on_ok())
    dialog.bind("<Escape>", lambda event: on_cancel())
    
    # 等待對話框關閉
    dialog.wait_window()
    
    return result[0]

def parse_word_document_com(self, filepath: str):
    """
    使用 Windows COM 與 Microsoft Word 互動來解析 .docx 文件，
    以嘗試獲取包括自動編號在內的渲染後文字。

    Args:
        filepath (str): Word 文件的路徑。

    Returns:
        str | None: 解析後的文字內容，包含自動編號和縮排。
                      如果發生錯誤、缺少依賴或無法執行則返回 None。
    """
    if not HAS_PYWIN32 or platform.system() != 'Windows':
        return None
    
    try:
        # 初始化 COM 環境
        pythoncom.CoInitialize()
        
        # 創建 Word 應用程序實例
        word_app = win32.Dispatch("Word.Application")
        word_app.Visible = False  # 不顯示 Word 視窗
        
        try:
            # 打開文檔
            doc = word_app.Documents.Open(filepath)
            
            try:
                # 使用更詳細的方法獲取文檔內容，保留格式和自動編號
                paragraphs = []
                
                # 迭代文件中的段落
                for i, para in enumerate(doc.Paragraphs):
                    try:
                        para_range = para.Range
                        list_string = para_range.ListFormat.ListString
                        full_text = para_range.Text.rstrip('\r\n')
                        
                        # 計算縮排
                        indent_space = ""
                        try:
                            indent_points = para.Format.LeftIndent
                            if indent_points > 0:
                                indent_level = int(indent_points / 18)  # 每18磅算一級縮排
                                indent_space = " " * (indent_level * 3)  # 每級縮排3個空格
                        except:
                            pass
                        
                        # 組合輸出，確保自動編號被保留
                        if list_string:
                            # 如果有自動編號，確保它被正確顯示
                            temp_text = full_text
                            if temp_text.startswith(list_string):
                                temp_text = temp_text[len(list_string):].lstrip()
                            formatted_line = f"{indent_space}{list_string} {temp_text}"
                        else:
                            formatted_line = f"{indent_space}{full_text}"
                        
                        paragraphs.append(formatted_line)
                    except Exception as para_exc:
                        print(f"讀取段落 {i+1} 時發生錯誤: {para_exc}")
                        paragraphs.append(f"[讀取段落錯誤]")
                
                # 使用換行符連接段落，保留格式
                content = '\n'.join(paragraphs)
                
                # 移除文本末尾可能出現的斜線
                if content.endswith('/'):
                    content = content[:-1]
                content = content.replace('/\n', '\n')
                content = content.replace('/ \n', '\n')
                
                return content
            finally:
                # 關閉文檔
                doc.Close(SaveChanges=False)
        finally:
            # 退出 Word 應用程序
            word_app.Quit()
            
            # 釋放 COM 環境
            pythoncom.CoUninitialize()
    
    except Exception as e:
        print(f"COM 解析失敗: {str(e)}")
        return None
