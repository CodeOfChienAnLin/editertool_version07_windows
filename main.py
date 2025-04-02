import os
import sys
import json
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, simpledialog
import threading
import docx2txt  # 用於讀取Word文檔
import msoffcrypto  # 用於處理加密的Office文檔
import io
from io import BytesIO
import opencc  # 用於中文文字轉換和校正
import tempfile
from docx import Document  # 用於更精確地讀取Word文檔格式
from PIL import Image, ImageTk
import datetime
import traceback
import logging
import platform # 用於檢查操作系統
from pathlib import Path # --- 將 Path 導入移到這裡 ---

# --- COM 相關匯入和檢查 (用於 Windows + Word 解析) ---
HAS_PYWIN32 = False
if platform.system() == 'Windows': # COM 只在 Windows 上可用
    try:
        import win32com.client as win32
        import pythoncom # 需要初始化 COM
        HAS_PYWIN32 = True
    except ImportError:
        print("警告：未找到 pywin32 模組。COM 功能需要 Windows、Microsoft Word 和 'pip install pywin32'。")
        pass # 即使沒有 pywin32，程式仍可嘗試啟動 (但功能受限)

class TextCorrectionTool:
    """文字校正工具主類別"""
    def __init__(self, root):
        """初始化應用程式

        參數:
            root: tkinter的根視窗
        """
        self.root = root
        self.root.title("編審神器")
        self.root.geometry("1000x700")  # 設定視窗大小為1000x700
        self.root.resizable(False, False)  # 禁止調整視窗大小

        # 設定錯誤日誌
        self.setup_error_logging()

        # 載入詞彙保護表
        self.protected_words = self.load_protected_words()

        # 載入設定
        self.settings = self.load_settings()

        # 初始化OpenCC轉換器
        try:
            # 使用簡體到繁體的轉換
            self.converter = opencc.OpenCC('s2t')  # 將簡體字轉為繁體字
        except Exception as e:
            messagebox.showerror("錯誤", f"無法初始化OpenCC轉換器: {str(e)}")
            self.converter = None

        self.create_widgets()  # 創建UI元件
        self.setup_drag_drop()  # 設置拖放功能

        # 圖片相關變數
        self.images = []  # 存儲原始圖片
        self.image_refs = []  # 存儲 Tkinter PhotoImage 引用
        self.download_path = os.path.join(os.path.expanduser("~"), "Downloads")  # 預設下載路徑

        # 應用深色模式設定
        self.apply_theme()

    def setup_error_logging(self):
        """設定錯誤日誌記錄"""
        # 確保日誌目錄存在
        log_dir = "logs"
        if not os.path.exists(log_dir):
            os.makedirs(log_dir)

        # 設定日誌檔案名稱（包含日期）
        log_file = os.path.join(log_dir, f"error_log_{datetime.datetime.now().strftime('%Y%m%d')}.log")

        # 配置日誌記錄器
        logging.basicConfig(
            filename=log_file,
            level=logging.ERROR,
            format='%(asctime)s - %(levelname)s - %(message)s',
            datefmt='%Y-%m-%d %H:%M:%S'
        )

        # 設定未捕獲異常的處理器
        def handle_exception(exc_type, exc_value, exc_traceback):
            """處理未捕獲的異常"""
            if issubclass(exc_type, KeyboardInterrupt):
                # 正常退出程式的情況，不記錄
                sys.__excepthook__(exc_type, exc_value, exc_traceback)
                return

            # 記錄詳細的錯誤信息
            error_msg = "".join(traceback.format_exception(exc_type, exc_value, exc_traceback))
            logging.error(f"未捕獲的異常:\n{error_msg}")

            # 顯示錯誤訊息給使用者
            messagebox.showerror("程式錯誤", f"發生嚴重錯誤，程式可能需要重新啟動。\n錯誤已記錄到日誌檔案中。\n\n錯誤類型: {exc_type.__name__}\n錯誤訊息: {str(exc_value)}")

        # 設定全局異常處理器
        sys.excepthook = handle_exception

    def create_widgets(self):
        """創建所有UI元件"""
        # 選單列
        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)

        # 檔案選單
        file_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="檔案", menu=file_menu)
        file_menu.add_command(label="開啟", command=self.open_file)
        file_menu.add_command(label="儲存", command=self.save_file)
        file_menu.add_separator()
        file_menu.add_command(label="離開", command=self.root.quit)

        # 編輯選單 (移除文字修正, 加入還原)
        edit_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="編輯", menu=edit_menu)
        edit_menu.add_command(label="還原上一步", command=self.undo_last_action) # 加入還原
        edit_menu.add_separator()
        # edit_menu.add_command(label="校正文字", command=self.correct_text) # 移至工具欄
        edit_menu.add_command(label="管理保護詞彙", command=self.manage_protected_words)
        edit_menu.add_command(label="清除紅色標記", command=self.clear_correction_highlights)


        # 設定選單
        settings_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="設定", menu=settings_menu)
        settings_menu.add_command(label="文字格式", command=self.open_text_settings)
        settings_menu.add_command(label="換色模式", command=self.toggle_dark_mode)

        # 檢視選單
        view_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="檢視", menu=view_menu)
        view_menu.add_command(label="錯誤日誌", command=self.view_error_logs)

        # 主框架
        main_frame = tk.Frame(self.root)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # 創建標籤頁控件
        self.notebook = ttk.Notebook(main_frame)
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # 創建文字修正標籤頁
        self.text_correction_tab = tk.Frame(self.notebook)
        self.notebook.add(self.text_correction_tab, text="文字修正")

        # 創建代辦事項標籤頁
        self.notes_tab = tk.Frame(self.notebook)
        self.notebook.add(self.notes_tab, text="代辦事項")
        # Add a label to the notes tab for visibility testing
        tk.Label(self.notes_tab, text="代辦事項功能區").pack(pady=20)


        # --- 新增工具欄框架 ---
        self.toolbar_main_frame = tk.Frame(self.text_correction_tab, relief=tk.RAISED, bd=1)
        self.toolbar_main_frame.pack(side=tk.TOP, fill=tk.X, padx=5, pady=(5, 0)) # Pack toolbar first

        # 工具欄上層
        self.toolbar_top_frame = tk.Frame(self.toolbar_main_frame)
        self.toolbar_top_frame.pack(side=tk.TOP, fill=tk.X)

        # 工具欄按鈕 (上層)
        self.undo_button = tk.Button(self.toolbar_top_frame, text="還原上一步", command=self.undo_last_action)
        self.undo_button.pack(side=tk.LEFT, padx=2, pady=2)

        self.correct_button = tk.Button(self.toolbar_top_frame, text="文字修正", command=self.correct_text)
        self.correct_button.pack(side=tk.LEFT, padx=2, pady=2)

        self.add_shortcut_button = tk.Button(self.toolbar_top_frame, text="新增快捷字", command=self.add_shortcut) # Placeholder command
        self.add_shortcut_button.pack(side=tk.LEFT, padx=2, pady=2)

        # 工具欄下層
        self.toolbar_bottom_frame = tk.Frame(self.toolbar_main_frame)
        self.toolbar_bottom_frame.pack(side=tk.TOP, fill=tk.X)

        # 工具欄按鈕 (下層 - 快捷字/符號)
        shortcuts = ["，", "。", "「」", "『』", "民國(下同)", "新臺幣(下同)"]
        for sc in shortcuts:
            # Handle quotes needing cursor placement inside
            if sc == "「」" or sc == "『』":
                btn = tk.Button(self.toolbar_bottom_frame, text=sc,
                                command=lambda s=sc: self.insert_text_at_cursor(s, move_cursor=True))
            else:
                btn = tk.Button(self.toolbar_bottom_frame, text=sc,
                                command=lambda s=sc: self.insert_text_at_cursor(s))
            btn.pack(side=tk.LEFT, padx=2, pady=2)


        # --- 圖片顯示區域框架 (Pack at the bottom) ---
        self.image_frame = tk.Frame(self.text_correction_tab, height=120, bg="white") # Fixed height
        self.image_frame.pack(side=tk.BOTTOM, fill=tk.X, padx=5, pady=5) # Pack image frame at the bottom
        self.image_frame.pack_propagate(False)

        # 圖片顯示區域的滾動畫布
        self.image_canvas = tk.Canvas(self.image_frame, bg="white")
        self.image_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # 圖片區域的垂直滾動條
        image_scrollbar = tk.Scrollbar(self.image_frame, orient=tk.VERTICAL, command=self.image_canvas.yview)
        image_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.image_canvas.configure(yscrollcommand=image_scrollbar.set)

        # 創建一個框架來放置圖片
        self.image_container = tk.Frame(self.image_canvas, bg="white")
        self.image_canvas.create_window((0, 0), window=self.image_container, anchor="nw")

        # 綁定圖片容器的配置事件
        self.image_container.bind("<Configure>", self.on_image_container_configure)

        # 按鈕框架 (圖片下載)
        img_button_frame = tk.Frame(self.image_frame, bg="white") # Renamed to avoid conflict
        img_button_frame.pack(side=tk.RIGHT, fill=tk.Y, padx=5, pady=5)

        # 下載圖片按鈕
        self.download_button = tk.Button(img_button_frame, text="下載圖片", command=self.download_images)
        self.download_button.pack(side=tk.TOP, fill=tk.X, padx=5, pady=2)

        # 選擇路徑按鈕
        self.path_button = tk.Button(img_button_frame, text="選擇路徑", command=self.choose_download_path)
        self.path_button.pack(side=tk.TOP, fill=tk.X, padx=5, pady=2)

        # --- 文字處理區域框架 (Pack last to fill remaining space) ---
        text_frame = tk.Frame(self.text_correction_tab) # Removed width
        text_frame.pack(side=tk.TOP, fill=tk.BOTH, expand=True, padx=5, pady=5) # Pack after toolbar, before image_frame

        # 添加垂直滾動條
        y_scrollbar = tk.Scrollbar(text_frame, orient=tk.VERTICAL)
        y_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        # 文字處理區域 - 啟用 undo
        self.text_area = tk.Text(text_frame,
                               font=(self.settings["font_family"], self.settings["font_size"]),
                               spacing3=self.settings["line_spacing"],
                               wrap=tk.WORD,
                               undo=True, # 啟用 Undo/Redo
                               yscrollcommand=y_scrollbar.set)
        self.text_area.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # 創建紅色底線標籤
        self.text_area.tag_configure("corrected", underline=True, underlinefg="red")

        # 設置縮進
        self.text_area.config(tabs=("1c", "2c", "3c", "4c"), tabstyle="wordprocessor")

        # 綁定事件
        self.text_area.bind("<<Modified>>", self.adjust_indentation)

        # 設置滾動條命令
        y_scrollbar.config(command=self.text_area.yview)

        # 狀態欄
        self.status_bar = tk.Label(self.root, text="就緒", bd=1, relief=tk.SUNKEN, anchor=tk.W)
        self.status_bar.pack(side=tk.BOTTOM, fill=tk.X)


    def undo_last_action(self):
        """還原上一步文字編輯操作"""
        try:
            self.text_area.edit_undo()
            self.status_bar.config(text="已還原上一步操作")
        except tk.TclError:
            self.status_bar.config(text="沒有可還原的操作")

 # --- 修改：實作 add_shortcut ---
    def add_shortcut(self):
        """新增快捷字: 跳出輸入視窗，將輸入文字變成新按鈕加到工具欄下層"""
        shortcut_text = simpledialog.askstring(
            "新增快捷字",
            "請輸入要新增的快捷字:",
            parent=self.root  # Make the dialog modal to the main window
        )

        if shortcut_text:  # Check if the user entered text (didn't cancel)
            shortcut_text = shortcut_text.strip() # Remove leading/trailing whitespace
            if shortcut_text: # Check again after stripping to ensure it's not just whitespace
                try:
                    # --- Create the new button ---
                    # IMPORTANT: Use lambda s=shortcut_text: ...
                    # This captures the *current* value of shortcut_text for this specific button's command.
                    # If you just used lambda: self.insert_text_at_cursor(shortcut_text),
                    # all dynamically created buttons would insert the *last* added shortcut text.
                    new_button = tk.Button(
                        self.toolbar_bottom_frame,
                        text=shortcut_text,
                        command=lambda s=shortcut_text: self.insert_text_at_cursor(s)
                    )

                    # --- Pack the new button ---
                    # It will automatically appear to the right of existing buttons packed with side=tk.LEFT
                    new_button.pack(side=tk.LEFT, padx=2, pady=2)

                    # --- Apply the current theme to the new button ---
                    # We need to ensure new buttons also follow the theme.
                    self.apply_theme_to_widget(new_button) # 使用輔助函數套用主題

                    self.status_bar.config(text=f"已新增快捷字: {shortcut_text}")
                    print(f"新增快捷字: {shortcut_text}")

                except Exception as e:
                    error_msg = f"無法新增快捷按鈕: {str(e)}"
                    self.log_error("Shortcut Error", error_msg, traceback.format_exc()) # 使用您的錯誤記錄
            else:
                 # User entered only whitespace
                 messagebox.showwarning("提示", "快捷字不能為空或僅包含空白字符。", parent=self.root)
        # else: User cancelled the dialog, do nothing.

    # --- 新增：套用主題到單一元件的輔助函數 ---
    def apply_theme_to_widget(self, widget):
         """Apply the current theme to a specific widget."""
         if not hasattr(self, 'settings') or not widget: # Safety check
             return

         # Determine colors based on theme
         if self.settings["dark_mode"]:
             bg_color = "#2b2b2b" # 背景色
             fg_color = "white" # 前景色 (文字)
             button_bg = "#3c3f41" # 按鈕背景
             button_fg = "white" # 按鈕文字
             # Add more specific colors if needed
         else:
             # 使用系統預設顏色以獲得更好的外觀整合
             bg_color = "SystemButtonFace" # 主背景色
             fg_color = "black" # 前景色
             button_bg = "SystemButtonFace" # 按鈕背景
             button_fg = "black" # 按鈕文字
             # Add more specific colors if needed

         # Apply theme based on widget type (expand as needed)
         widget_type = widget.winfo_class() # 獲取元件類型 (例如 'Button', 'Label', 'Frame')

         try:
             if widget_type == 'Button':
                 widget.configure(bg=button_bg, fg=button_fg)
             elif widget_type == 'Label':
                 # 確保標籤背景與其容器一致
                 # Frame 的背景色通常是 bg_color，所以標籤也用 bg_color
                  widget.configure(bg=bg_color, fg=fg_color)
             elif widget_type == 'Frame':
                 widget.configure(bg=bg_color) # Frame 本身使用主背景色
             elif widget_type == 'Canvas': # 例如圖片區的 Canvas
                 # Canvas 的背景可能需要特別設定 (例如淺色模式下的白色)
                 canvas_bg = "white" if not self.settings["dark_mode"] else "#2b2b2b"
                 widget.configure(bg=canvas_bg)
             # Add other common widget types if necessary
         except tk.TclError:
              # Widget might already be destroyed, ignore error
              pass

    # --- 修改：更新 apply_theme 以使用輔助函數並涵蓋更多元件 ---
    def apply_theme(self):
        """應用主題設定"""
        if self.settings["dark_mode"]:
            # 深色模式
            bg_color = "#2b2b2b"
            fg_color = "white"
            text_bg = "#2b2b2b" # 文字區域背景
            text_fg = "white" # 文字區域文字
            button_bg = "#3c3f41"
            button_fg = "white"
            canvas_bg = "#2b2b2b" # Canvas 背景 (圖片區)
            toolbar_bg = "#3c3f41" # 工具列背景
            img_container_bg = "#2b2b2b" # 圖片容器 Frame 背景
            cursor_color = "white" # 文字區域游標顏色
            notebook_bg = "#2b2b2b" # Notebook 背景
            tab_bg = "#3c3f41" # Tab 背景
            selected_tab_bg = "#4f5254" # 選中 Tab 背景 (可調整)
        else:
            # 淺色模式 (使用系統預設)
            bg_color = "SystemButtonFace"
            fg_color = "black"
            text_bg = "white"
            text_fg = "black"
            button_bg = "SystemButtonFace"
            button_fg = "black"
            canvas_bg = "white"
            toolbar_bg = "SystemButtonFace"
            img_container_bg = "white"
            cursor_color = "black"
            notebook_bg = "SystemButtonFace"
            tab_bg = "SystemButtonFace"
            selected_tab_bg = "SystemHighlight" # 使用系統高亮色 (或自訂淺灰色)


        # 應用主題到主視窗
        self.root.configure(bg=bg_color)

        # 應用主題到文字區域
        # 使用 insertbackground 讓游標在深色模式下可見
        self.text_area.configure(bg=text_bg, fg=text_fg, insertbackground=cursor_color)

        # 應用主題到圖片區域 Frame
        self.image_frame.configure(bg=bg_color)
        # 套用主題到 Canvas 本身
        self.image_canvas.configure(bg=canvas_bg)
        # 套用主題到 Canvas *內部* 的容器 Frame
        self.image_container.configure(bg=img_container_bg)
        # 套用主題到容器內的圖片標籤 (假設它們是 Label)
        for child in self.image_container.winfo_children():
             self.apply_theme_to_widget(child) # 使用輔助函數

        # 應用主題到圖片下載按鈕框架及其按鈕
        img_button_frame = None
        # 穩健地找到包含下載/路徑按鈕的框架
        # (假設它被 pack 在 image_frame 的右側)
        for widget in self.image_frame.winfo_children():
            if isinstance(widget, tk.Frame) and widget.winfo_manager() == 'pack':
                pack_info = widget.pack_info()
                if 'side' in pack_info and pack_info['side'] == tk.RIGHT:
                    img_button_frame = widget
                    break

        if img_button_frame:
             img_button_frame.configure(bg=bg_color) # 設定框架本身的背景
             for child in img_button_frame.winfo_children():
                 self.apply_theme_to_widget(child) # 設定內部按鈕的主題


        # 應用主題到狀態欄
        self.status_bar.configure(bg=bg_color, fg=fg_color)

        # 應用主題到工具欄 (使用輔助函數以保持一致性)
        if hasattr(self, 'toolbar_main_frame'):
            self.toolbar_main_frame.configure(bg=toolbar_bg)
            self.toolbar_top_frame.configure(bg=toolbar_bg)
            self.toolbar_bottom_frame.configure(bg=toolbar_bg)
            # 套用主題到兩個工具列的所有子元件
            for child in self.toolbar_top_frame.winfo_children():
                self.apply_theme_to_widget(child)
            for child in self.toolbar_bottom_frame.winfo_children(): # 包括動態新增的按鈕
                self.apply_theme_to_widget(child)

        # 應用主題到 Notebook 標籤頁 (使用 ttk.Style)
        style = ttk.Style()
        # 設定 Notebook 元件本身的背景 (標籤周圍的區域)
        style.configure("TNotebook", background=bg_color)
        # 設定 Notebook 內部頁面 (Frame) 的背景
        style.configure("TFrame", background=bg_color) # 確保 Tab 內的 Frame 背景也套用
        # 設定 Tab 的樣式
        style.configure("TNotebook.Tab",
                        background=tab_bg,         # 未選中 Tab 的背景
                        foreground=fg_color,       # 未選中 Tab 的文字顏色
                        padding=[5, 2]             # 增加一些內邊距
                       )
        # 設定選中和活動狀態下 Tab 的樣式
        style.map("TNotebook.Tab",
                  background=[("selected", selected_tab_bg)], # 選中 Tab 的背景
                  foreground=[("selected", fg_color)]        # 選中 Tab 的文字顏色
                 )

        # 應用主題到標籤頁內部的 Frame (確保它們有背景色)
        self.text_correction_tab.configure(bg=bg_color)
        self.notes_tab.configure(bg=bg_color)
        # 套用主題到 notes_tab 內的元件 (如果有)
        for child in self.notes_tab.winfo_children():
             self.apply_theme_to_widget(child) # 例如: 套用到標籤頁內部的 Frame

    def insert_text_at_cursor(self, text_to_insert, move_cursor=False):
        """在目前游標位置插入文字"""
        try:
            insert_index = self.text_area.index(tk.INSERT)
            self.text_area.insert(insert_index, text_to_insert)
            if move_cursor and len(text_to_insert) > 1:
                # Move cursor back one position (typically for quotes)
                # Calculate the new index based on the inserted text length
                line, col = map(int, str(insert_index).split('.'))
                new_col = col + len(text_to_insert) - 1
                new_index = f"{line}.{new_col}"
                self.text_area.mark_set(tk.INSERT, new_index)
            self.text_area.focus_set() # Keep focus on text area
            # Manually trigger Modified event for undo stack if needed
            self.text_area.edit_modified(True)
        except Exception as e:
            print(f"插入文字時出錯: {e}")
            self.status_bar.config(text=f"插入文字時出錯: {e}")


    def on_image_container_configure(self, event):
        """當圖片容器大小變化時，更新畫布的滾動區域"""
        self.image_canvas.configure(scrollregion=self.image_canvas.bbox("all"))

    def setup_drag_drop(self):
        """設置拖放功能"""
        try:
            # 直接使用 Tkinter 原生的拖放功能
            # 為文字區域啟用拖放
            self.text_area.drop_target_register('DND_Files')
            self.text_area.dnd_bind('<<Drop>>', self.handle_drop)
            self.status_bar.config(text="拖放功能已啟用，可以拖放 Word 文檔")
            print("已啟用 Tkinter 原生拖放功能")
        except Exception as e:
            print(f"Tkinter 原生拖放初始化失敗: {str(e)}")

            # 嘗試使用 TkDND
            try:
                print("嘗試使用 TkDND...")
                # 嘗試將 TkDND 套件目錄加入路徑
                import sys
                import os
                tkdnd_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), 'tkdnd'))
                if os.path.exists(tkdnd_dir):
                    sys.path.append(tkdnd_dir)

                from tkinter import TkVersion
                if TkVersion >= 8.6:
                    # Tk 8.6+ 原生支援拖放
                    self.root.tk.call('package', 'require', 'tkdnd')
                    self.root.tk.call('tkdnd::drop_target', 'register', self.text_area._w)
                    self.root.tk.call('bind', self.text_area._w, '<<Drop>>',
                                     self.root.register(self.handle_drop))
                    print("使用 Tk 8.6+ 原生拖放功能")
                    self.status_bar.config(text="拖放功能已啟用，可以拖放 Word 文檔")
                    return

                # 嘗試使用 TkinterDnD2
                try:
                    print("嘗試使用 TkinterDnD2...")
                    # 使用絕對導入
                    from TkinterDnD2 import TkinterDnD, DND_FILES
                    TkinterDnD.dnd_start(self.root)
                    self.text_area.drop_target_register(DND_FILES)
                    self.text_area.dnd_bind('<<Drop>>', self.handle_drop)
                    print("使用 TkinterDnD2 拖放功能")
                    self.status_bar.config(text="拖放功能已啟用，可以拖放 Word 文檔")
                    return
                except Exception as e:
                    print(f"TkinterDnD2 初始化失敗: {str(e)}")

                # 嘗試使用自訂的 TkDND 包裝類
                try:
                    print("嘗試使用自訂 TkDND 包裝類...")
                    from tkdnd_wrapper import TkDND
                    dnd = TkDND(self.root)
                    success = dnd.bindtarget(self.text_area, self.handle_drop, 'text/uri-list')
                    if success:
                        print("使用自訂 TkDND 包裝類")
                        self.status_bar.config(text="拖放功能已啟用，可以拖放 Word 文檔")
                        return
                except Exception as e:
                    print(f"自訂 TkDND 包裝類初始化失敗: {str(e)}")

            except Exception as e:
                print(f"TkDND 相關初始化失敗: {str(e)}")

            # 最後嘗試使用簡易的捕獲方法
            try:
                print("嘗試使用簡易捕獲方法...")
                # 處理貼上事件
                self.root.bind("<FocusIn>", self.check_clipboard)
                self.root.bind("<ButtonRelease>", self.check_clipboard)
                self.root.bind("<Key>", lambda e: self.check_clipboard() if e.keysym == 'v' and (e.state & 4) else None)
                print("已啟用簡易捕獲方法")
                self.status_bar.config(text="已啟用替代拖放功能，將檔案拖放到視窗後請點擊")
                return
            except Exception as e:
                print(f"簡易捕獲方法初始化失敗: {str(e)}")

            # 所有方法都失敗
            print("所有拖放方法都失敗了")
            self.status_bar.config(text="拖放功能初始化失敗，請使用選單開啟檔案")
            messagebox.showwarning("拖放功能警告", "拖放功能無法初始化\n請使用選單開啟檔案")

    def check_clipboard(self, event=None):
        """檢查剪貼簿是否有檔案路徑"""
        try:
            clipboard = self.root.clipboard_get()
            if clipboard and os.path.exists(clipboard) and clipboard.lower().endswith(('.docx', '.doc')):
                print(f"從剪貼簿獲取檔案: {clipboard}")
                self.process_word_file(clipboard)
                return True
        except Exception as e:
            print(f"檢查剪貼簿時發生錯誤: {str(e)}")
        return False

    def handle_drop(self, event):
        """處理檔案拖放事件

        參數:
            event: 拖放事件物件
        """
        try:
            data = event.data
            file_path = str(data).strip()

            print(f"原始拖放路徑: {file_path}")

            # 處理可能的格式
            # Windows 可能會在路徑周圍添加大括號或引號
            if file_path.startswith('{') and file_path.endswith('}'):
                file_path = file_path[1:-1]

            # 移除可能的引號
            if (file_path.startswith('"') and file_path.endswith('"')) or \
               (file_path.startswith("'") and file_path.endswith("'")):
                file_path = file_path[1:-1]

            # 處理可能的檔案URL格式
            if file_path.startswith('file:///'):
                file_path = file_path[8:].replace('/', '\\')

            # 處理 Mac 路徑格式或其他非標準路徑
            if file_path.startswith('/Mac/') or '://' in file_path:
                # 嘗試從路徑中提取實際的文件名
                file_name = os.path.basename(file_path)

                # 顯示錯誤訊息
                messagebox.showinfo("路徑格式不支援",
                                   f"檢測到非標準路徑格式: {file_path}\n\n"
                                   f"請嘗試以下方法：\n"
                                   f"1. 使用「檔案」選單中的「開啟」功能\n"
                                   f"2. 從檔案總管直接拖放檔案\n"
                                   f"3. 確保檔案位於本機上，而非網路位置")
                return

            print(f"處理後的檔案路徑: {file_path}")

            # 檢查檔案是否存在
            if not os.path.exists(file_path):
                messagebox.showerror("錯誤", f"找不到檔案: {file_path}\n請確保檔案路徑正確且檔案存在。")
                return

            # 檢查檔案是否為Word檔案
            if not file_path.lower().endswith(('.doc', '.docx')):
                messagebox.showerror("錯誤", f"不支援的檔案格式: {file_path}\n僅支援 .doc 和 .docx 格式。")
                return

            # 更新狀態欄
            self.status_bar.config(text=f"正在處理檔案: {os.path.basename(file_path)}")

            # 嘗試處理Word檔案
            try:
                # 先嘗試檢查文件是否加密
                try:
                    with open(file_path, 'rb') as f:
                        try:
                            office_file = msoffcrypto.OfficeFile(f)
                            if office_file.is_encrypted():
                                print("檔案已加密，需要密碼")
                                # 文件已加密，直接調用密碼處理方法
                                self.handle_password_protected_file(file_path)
                                return
                        except Exception as e:
                            print(f"檢查加密狀態時發生錯誤: {str(e)}")
                            # 繼續嘗試普通處理
                except Exception as e:
                    print(f"開啟檔案時發生錯誤: {str(e)}")
                            # 繼續嘗試普通處理
                except Exception as e:
                    print(f"開啟檔案時發生錯誤: {str(e)}")
                    # 繼續嘗試普通處理

                # --- 修改：調用新的處理邏輯 ---
                self.load_and_display_word_content(file_path)

            except Exception as e:
                # 檢查是否為加密文件的錯誤
                error_str = str(e).lower()
                if self._is_password_error(error_str):
                    # 可能是加密文件，嘗試使用密碼處理
                    print(f"檢測到加密錯誤: {error_str}")
                    self.handle_password_protected_file(file_path)
                else:
                    # 其他錯誤，顯示錯誤訊息
                    error_msg = f"處理檔案時發生錯誤: {str(e)}\n{traceback.format_exc()}"
                    print(error_msg)
                    logging.error(error_msg) # 記錄錯誤
                    messagebox.showerror("錯誤", f"處理檔案時發生錯誤: {str(e)}")
                    self.status_bar.config(text=f"處理檔案時發生錯誤") # 簡化狀態欄訊息

        except Exception as e:
            error_msg = f"處理拖放檔案時發生嚴重錯誤: {str(e)}\n{traceback.format_exc()}"
            print(error_msg)
            logging.error(error_msg) # 記錄錯誤
            self.status_bar.config(text=f"處理拖放檔案時發生錯誤")
            messagebox.showerror("錯誤", f"處理拖放檔案時發生嚴重錯誤: {str(e)}")

    def load_and_display_word_content(self, file_path, password=None):
        """
        載入並顯示 Word 文件內容，整合 COM 和原有方法，並提取圖片。

        參數:
            file_path: Word檔案路徑
            password: 檔案密碼（如果有的話）
        """
        self.status_bar.config(text=f"正在處理檔案: {os.path.basename(file_path)}...")
        self.root.update_idletasks() # 更新UI以顯示狀態

        # 清空之前的圖片
        self.clear_images()

        text = None
        method_used = "None"

        try:
            # --- 優先嘗試 COM 解析 (Windows 且有 pywin32) ---
            if platform.system() == 'Windows' and HAS_PYWIN32 and not password: # COM 方法通常不直接處理加密文件
                print(f"嘗試使用 COM 方法解析: {file_path}")
                self.status_bar.config(text=f"嘗試使用 COM 方法解析...")
                self.root.update_idletasks()
                text = self.parse_word_document_com(file_path)
                if text is not None:
                    method_used = "COM (MS Word)"
                    print("COM 解析成功")
                else:
                    print("COM 解析失敗，嘗試回退...")
                    self.status_bar.config(text=f"COM 解析失敗，嘗試其他方法...")
                    self.root.update_idletasks()

            # --- 如果 COM 失敗、不可用或需要密碼，則使用原有方法 ---
            if text is None:
                print(f"嘗試使用內建方法解析 (或處理加密文件): {file_path}")
                self.status_bar.config(text=f"嘗試使用內建方法解析...")
                self.root.update_idletasks()
                text = self.process_word_file_internal(file_path, password)
                if text is not None:
                    method_used = "內建 (docx2txt/python-docx)"
                    print("內建方法解析成功")
                else:
                    print("內建方法解析也失敗")

            # --- 如果成功獲取文本，則提取圖片 ---
            if text is not None:
                print("文本解析成功，開始提取圖片...")
                self.status_bar.config(text=f"文本解析成功 ({method_used})，正在提取圖片...")
                self.root.update_idletasks()
                # 圖片提取總是使用 python-docx，即使文本是用 COM 解析的
                # 如果文件已解密到臨時文件，需要傳遞臨時文件路徑
                file_to_extract_images_from = file_path
                if password:
                    # 如果有密碼，process_word_file_internal 會處理解密和臨時文件
                    # 我們需要獲取那個臨時文件的路徑來提取圖片
                    # (這部分需要在 process_word_file_internal 中返回臨時路徑或修改邏輯)
                    # 暫時假設 process_word_file_internal 會處理好圖片提取
                    # 或者，我們可以在這裡重新解密一次專門用於提圖，但效率低
                    # **目前的 process_word_file_internal 結構會在解密後提取圖片，所以這裡可能不需要額外操作**
                    # **但需要確認 process_word_file_internal 的圖片提取邏輯是否正確觸發**
                    # **修改 process_word_file_internal，使其在解密後調用 extract_images_from_docx**
                    pass # 假設 process_word_file_internal 已處理圖片提取
                else:
                    # 對於無密碼文件，直接提取
                    self.extract_images_from_docx(file_to_extract_images_from)

                # --- 更新 UI ---
                self.text_area.delete(1.0, tk.END)
                self.text_area.insert(tk.END, text)
                self.status_bar.config(text=f"已載入: {os.path.basename(file_path)} (使用 {method_used})")
                self.adjust_indentation()
                self.correct_text()

            else:
                # 所有方法都失敗
                messagebox.showerror("錯誤", f"無法讀取檔案 '{os.path.basename(file_path)}'。\n兩種解析方法均失敗。")
                self.status_bar.config(text=f"讀取檔案失敗: {os.path.basename(file_path)}")

        except FileNotFoundError as fnf_error:
             messagebox.showerror("錯誤", f"找不到檔案: {fnf_error}")
             self.status_bar.config(text="找不到檔案")
        except msoffcrypto.exceptions.DecryptionError as decrypt_error:
             messagebox.showerror("解密錯誤", f"解密失敗，密碼可能不正確: {decrypt_error}")
             self.status_bar.config(text="解密失敗，密碼錯誤")
        except Exception as e:
            error_msg = f"處理 Word 文件時發生未知錯誤: {str(e)}\n{traceback.format_exc()}"
            print(error_msg)
            logging.error(error_msg)
            messagebox.showerror("錯誤", f"處理 Word 文件時發生未知錯誤: {str(e)}")
            self.status_bar.config(text="處理檔案時發生錯誤")


    def process_word_file_internal(self, file_path, password=None):
        """
        處理Word檔案（內部方法，優先使用 docx2txt，回退到 python-docx）。
        如果提供密碼，則先解密。
        **修改：** 在解密成功後調用圖片提取。

        參數:
            file_path: Word檔案路徑
            password: 檔案密碼（如果有的話）

        回傳:
            檔案內容 (str) 或 None (如果失敗)
        """
        # 檢查檔案是否存在
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"找不到檔案: {file_path}")

        # 如果提供了密碼，嘗試解密檔案
        if password:
            temp_path = None # 初始化
            try:
                # 創建一個臨時檔案來存儲解密後的內容
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                    temp_path = temp_file.name

                # 打開加密檔案
                with open(file_path, 'rb') as f:
                    file_bytes = f.read()

                # 創建一個 BytesIO 對象
                file_stream = BytesIO(file_bytes)

                # 使用 msoffcrypto 解密
                ms_file = msoffcrypto.OfficeFile(file_stream)
                ms_file.load_key(password=password)

                with open(temp_path, 'wb') as f:
                    ms_file.decrypt(f)

                # 處理解密後的檔案
                text = self._process_unencrypted_file(temp_path)

                # --- 新增：在解密成功後提取圖片 ---
                if text is not None:
                    print(f"解密成功，從臨時文件提取圖片: {temp_path}")
                    self.extract_images_from_docx(temp_path)
                else:
                    print("解密後文件文本提取失敗，跳過圖片提取")

                return text # 返回文本內容
            except msoffcrypto.exceptions.DecryptionError as decrypt_error:
                 print(f"解密失敗: {decrypt_error}")
                 raise # 重新拋出解密錯誤，讓上層處理
            except Exception as e:
                # 如果解密或後續處理失敗，拋出異常
                print(f"處理加密文件時出錯: {e}")
                raise Exception(f"處理加密文件失敗: {str(e)}")
            finally:
                 # 確保刪除臨時檔案
                 if temp_path and os.path.exists(temp_path):
                     try:
                         os.unlink(temp_path)
                         print(f"已刪除臨時文件: {temp_path}")
                     except OSError as unlink_error:
                         print(f"警告：無法刪除臨時文件 {temp_path}: {unlink_error}")
        else:
            # 處理未加密的檔案
            text = self._process_unencrypted_file(file_path)
            # --- 對於未加密文件，圖片提取由 load_and_display_word_content 調用 ---
            # self.extract_images_from_docx(file_path) # 不在這裡提取，由上層調用
            return text

    # process_word_file 函數現在被 load_and_display_word_content 和 process_word_file_internal 取代
    # 可以刪除舊的 process_word_file 函數，或者保留它作為內部實現細節（但不直接從 UI 調用）
    # 這裡選擇重命名舊函數為 internal，並讓 load_and_display_word_content 成為主要入口

    # def process_word_file(self, file_path, password=None):
    #     """(舊函數 - 不再直接使用) 處理Word檔案..."""
    #     # ... 舊的實現 ...
    #     pass


    def _process_unencrypted_file(self, file_path):
        """處理未加密的Word檔案

        參數:
            file_path: Word檔案路徑

        回傳:
            檔案內容
        """
        try:
            # 嘗試使用 docx2txt 提取文字
            try:
                import docx2txt
                text = docx2txt.process(file_path)
                print("使用 docx2txt 成功提取文字")
                return text
            except Exception as docx2txt_error:
                print(f"docx2txt 失敗: {docx2txt_error}，嘗試使用 python-docx")
                # 如果 docx2txt 失敗，嘗試使用 python-docx
                try:
                    import docx
                    doc = docx.Document(file_path)
                    text = "\n".join([para.text for para in doc.paragraphs])
                    print("使用 python-docx 成功提取文字")
                    return text
                except Exception as docx_error:
                    print(f"python-docx 失敗: {docx_error}，嘗試使用 COM")
                    # 如果 python-docx 也失敗，嘗試使用 COM 方法
                    text = self.parse_word_document_com(file_path)
                    if text:
                        print("使用 COM 成功提取文字")
                        return text
                    else:
                        raise Exception("所有提取方法都失敗")
        except Exception as e:
            # 檢查是否為加密文件的錯誤
            error_str = str(e).lower()
            if self._is_password_error(error_str):
                # 可能是加密文件，嘗試使用密碼處理
                print(f"檢測到加密錯誤: {error_str}")
                self.handle_password_protected_file(file_path)
            else:
                # 其他錯誤，顯示錯誤訊息
                messagebox.showerror("錯誤", f"處理檔案時發生錯誤: {str(e)}")
                self.status_bar.config(text=f"處理檔案時發生錯誤: {str(e)}")
            raise e  # 重新拋出異常，讓上層函數知道處理失敗

    def process_word_file(self, file_path, password=None):
        """處理Word檔案

        參數:
            file_path: Word檔案路徑
            password: 檔案密碼（如果有的話）

        回傳:
            檔案內容
        """
        # 檢查檔案是否存在
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"找不到檔案: {file_path}")

        # 清空之前的圖片
        self.clear_images()

        # 如果提供了密碼，嘗試解密檔案
        if password:
            try:
                # 創建一個臨時檔案來存儲解密後的內容
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                    temp_path = temp_file.name

                # 打開加密檔案
                with open(file_path, 'rb') as f:
                    file_bytes = f.read()

                # 創建一個 BytesIO 對象
                file_stream = BytesIO(file_bytes)

                # 使用 msoffcrypto 解密
                ms_file = msoffcrypto.OfficeFile(file_stream)
                ms_file.load_key(password=password)

                with open(temp_path, 'wb') as f:
                    ms_file.decrypt(f)

                # 處理解密後的檔案
                text = self._process_unencrypted_file(temp_path)

                # 提取圖片
                self.extract_images_from_docx(temp_path)

                # 刪除臨時檔案
                os.unlink(temp_path)

                return text
            except Exception as e:
                # 如果解密失敗，拋出異常
                raise Exception(f"解密失敗: {str(e)}")
        else:
            # 處理未加密的檔案
            text = self._process_unencrypted_file(file_path)

            # 提取圖片
            self.extract_images_from_docx(file_path)

            return text

    def _is_password_error(self, error_message):
        """檢查錯誤訊息是否與密碼保護相關

        參數:
            error_message: 錯誤訊息

        回傳:
            是否為密碼相關錯誤
        """
        error_message = error_message.lower()
        password_keywords = ["password", "encrypted", "保護", "密碼", "加密"]
        return any(keyword in error_message for keyword in password_keywords)

    def _extract_text_from_document(self, doc):
        """從 python-docx Document 物件中提取文字

        參數:
            doc: python-docx Document 物件

        回傳:
            提取的文字
        """
        # 提取文本，保留段落格式
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():  # 忽略空段落
                paragraphs.append(para.text)

        # 提取表格內容
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    paragraphs.append('\t'.join(row_text))

        # 使用兩個換行符連接段落，保留格式
        return '\n\n'.join(paragraphs)

    def extract_images_from_docx(self, file_path):
        """從Word文件中提取圖片

        參數:
            file_path: Word檔案路徑
        """
        try:
            # 使用 python-docx 打開文件
            doc = Document(file_path)

            # 提取文檔中的所有圖片
            image_index = 0
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        # 獲取圖片數據
                        image_data = rel.target_part.blob

                        # 使用 PIL 處理圖片
                        image = Image.open(BytesIO(image_data))

                        # 保存圖片到列表中
                        self.images.append(image)

                        # 顯示圖片
                        self.display_image(image, image_index)

                        image_index += 1
                    except Exception as e:
                        print(f"提取圖片時出錯: {str(e)}")

            # 更新狀態欄
            if image_index > 0:
                self.status_bar.config(text=f"已提取 {image_index} 張圖片")

        except Exception as e:
            print(f"提取圖片時出錯: {str(e)}")
            messagebox.showwarning("警告", f"提取圖片時出錯: {str(e)}")

    def display_image(self, image, index):
        """在圖片區域顯示圖片

        參數:
            image: PIL Image 對象
            index: 圖片索引
        """
        # 計算縮放後的圖片大小，最大高度為 100 像素
        max_height = 100
        width, height = image.size

        if height > max_height:
            ratio = max_height / height
            new_height = max_height
            new_width = int(width * ratio)
        else:
            new_width = width
            new_height = height

        # 縮放圖片
        resized_image = image.resize((new_width, new_height), Image.LANCZOS)

        # 轉換為 Tkinter 可用的格式
        tk_image = ImageTk.PhotoImage(resized_image)

        # 保存引用，防止垃圾回收
        self.image_refs.append(tk_image)

        # 創建標籤來顯示圖片
        image_label = tk.Label(self.image_container, image=tk_image, bg="white")
        image_label.grid(row=0, column=index, padx=5, pady=5, sticky="w")

        # 綁定點擊事件，以便放大查看
        image_label.bind("<Button-1>", lambda event, img=image, idx=index: self.show_full_image(img, idx))

    def show_full_image(self, image, index):
        """顯示原始大小的圖片

        參數:
            image: PIL Image 對象
            index: 圖片索引
        """
        # 創建新視窗
        image_window = tk.Toplevel(self.root)
        image_window.title(f"圖片 {index + 1}")

        # 獲取圖片原始大小
        width, height = image.size

        # 限制最大顯示尺寸
        max_width = 800
        max_height = 600

        if width > max_width or height > max_height:
            # 計算縮放比例
            width_ratio = max_width / width if width > max_width else 1
            height_ratio = max_height / height if height > max_height else 1
            ratio = min(width_ratio, height_ratio)

            # 縮放圖片
            new_width = int(width * ratio)
            new_height = int(height * ratio)
            display_image = image.resize((new_width, new_height), Image.LANCZOS)
        else:
            display_image = image
            new_width = width
            new_height = height

        # 設置視窗大小
        window_width = new_width + 20
        window_height = new_height + 20

        # 居中顯示
        screen_width = image_window.winfo_screenwidth()
        screen_height = image_window.winfo_screenheight()
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2
        image_window.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 轉換為 Tkinter 可用的格式
        tk_image = ImageTk.PhotoImage(display_image)

        # 創建標籤來顯示圖片
        image_label = tk.Label(image_window, image=tk_image)
        image_label.image = tk_image  # 保存引用
        image_label.pack(padx=10, pady=10)

        # 添加關閉按鈕
        close_button = tk.Button(image_window, text="關閉", command=image_window.destroy)
        close_button.pack(pady=5)

    def clear_images(self):
        """清空圖片區域"""
        # 清空圖片列表
        self.images = []
        self.image_refs = []

        # 清空圖片容器
        for widget in self.image_container.winfo_children():
            widget.destroy()

    def download_images(self):
        """下載所有圖片到指定路徑"""
        if not self.images:
            messagebox.showinfo("提示", "沒有可下載的圖片")
            return

        try:
            # 確保下載路徑存在
            os.makedirs(self.download_path, exist_ok=True)

            # 下載所有圖片
            for i, image in enumerate(self.images):
                # 生成檔案名稱
                file_name = f"image_{i + 1}.png"
                file_path = os.path.join(self.download_path, file_name)

                # 保存圖片
                image.save(file_path)

            # 更新狀態欄
            self.status_bar.config(text=f"已下載 {len(self.images)} 張圖片到 {self.download_path}")

            # 顯示成功訊息
            messagebox.showinfo("成功", f"已下載 {len(self.images)} 張圖片到:\n{self.download_path}")

        except Exception as e:
            messagebox.showerror("錯誤", f"下載圖片時出錯: {str(e)}")

    def choose_download_path(self):
        """選擇圖片下載路徑"""
        path = filedialog.askdirectory(title="選擇圖片下載路徑")
        if path:
            self.download_path = path
            self.status_bar.config(text=f"已設置下載路徑: {path}")

    def open_file(self):
        """開啟檔案對話框"""
        try:
            file_path = filedialog.askopenfilename(
                title="選擇Word檔案",
                filetypes=[("Word文件", "*.docx;*.doc"), ("所有檔案", "*.*")]
            )

            if file_path:
                print(f"選擇的檔案: {file_path}")
                # --- 修改：調用新的處理邏輯 ---
                self.load_and_display_word_content(file_path)

        except Exception as e:
            error_msg = f"開啟檔案時發生錯誤: {str(e)}\n{traceback.format_exc()}"
            print(error_msg)
            logging.error(error_msg) # 記錄錯誤
            self.status_bar.config(text="開啟檔案時出錯")
            messagebox.showerror("錯誤", f"無法開啟檔案: {str(e)}")

    def save_file(self):
        """儲存檔案對話框"""
        # 儲存檔案對話框
        file_path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("文字檔案", "*.txt"), ("Word文檔", "*.docx"), ("所有檔案", "*.*")]
        )
        if file_path:
            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    text = self.text_area.get(1.0, tk.END)
                    f.write(text)
                self.status_bar.config(text=f"已儲存到: {os.path.basename(file_path)}")
            except Exception as e:
                self.status_bar.config(text="儲存檔案時出錯")
                messagebox.showerror("錯誤", f"無法儲存檔案: {str(e)}")

    def correct_text(self):
        """校正文字內容"""
        # 檢查OpenCC轉換器是否正確初始化
        if not self.converter:
            self.status_bar.config(text="OpenCC轉換器未正確初始化，無法進行校正")
            messagebox.showerror("錯誤", "OpenCC轉換器未正確初始化，無法進行校正")
            return

        # 獲取文字內容
        text = self.text_area.get(1.0, tk.END)

        # 在背景執行校正，避免UI凍結
        threading.Thread(target=self._correct_text_thread, args=(text,)).start()

    def _correct_text_thread(self, text):
        """在背景執行文字校正的執行緒

        參數:
            text: 要校正的文字
        """
        try:
            print("開始文字校正執行緒")

            # 載入保護詞彙
            protected_words = self.load_protected_words()
            print(f"已載入保護詞彙: {protected_words}")

            # 創建一個暫時字典來保存保護詞彙的位置
            protected_positions = {}

            # 找出所有保護詞彙在文本中的位置
            for word in protected_words:
                start_pos = 0
                while True:
                    pos = text.find(word, start_pos)
                    if pos == -1:
                        break
                    protected_positions[pos] = (pos + len(word), word)
                    start_pos = pos + 1

            print(f"找到 {len(protected_positions)} 個保護詞彙位置")

            # 用於追蹤修改的部分
            corrections = []

            # 如果沒有保護詞彙，直接轉換整個文本
            if not protected_positions:
                original_text = text
                corrected_text = self.converter.convert(text)

                # 比較原始文本和校正後的文本，找出差異
                self._find_differences(original_text, corrected_text, corrections)
            else:
                # 分段處理文本，保護特定詞彙
                result = []
                last_end = 0

                # 按位置排序保護區域
                positions = sorted(protected_positions.keys())

                for start in positions:
                    end, word = protected_positions[start]

                    # 轉換保護詞彙前的文本
                    if start > last_end:
                        segment = text[last_end:start]
                        original_segment = segment
                        corrected_segment = self.converter.convert(segment)

                        # 比較原始文本和校正後的文本，找出差異
                        offset = last_end
                        self._find_differences(original_segment, corrected_segment, corrections, offset)

                        result.append(corrected_segment)

                    # 添加保護詞彙（不轉換）
                    result.append(word)
                    last_end = end

                # 處理最後一個保護詞彙之後的文本
                if last_end < len(text):
                    segment = text[last_end:]
                    original_segment = segment
                    corrected_segment = self.converter.convert(segment)

                    # 比較原始文本和校正後的文本，找出差異
                    offset = last_end
                    self._find_differences(original_segment, corrected_segment, corrections, offset)

                    result.append(corrected_segment)

                corrected_text = ''.join(result)

            print(f"校正完成，轉換後文字長度: {len(corrected_text)}")
            print(f"找到 {len(corrections)} 處修正")

            # 更新UI必須在主執行緒中進行
            self.root.after(0, self._update_text_area, corrected_text, corrections)
        except Exception as e:
            print(f"校正文字時發生錯誤: {str(e)}")
            # 更新UI必須在主執行緒中進行
            self.root.after(0, lambda: self.status_bar.config(text=f"校正文字時發生錯誤: {str(e)}"))
            self.root.after(0, lambda: messagebox.showerror("錯誤", f"校正文字時發生錯誤: {str(e)}"))

    def _find_differences(self, original_text, corrected_text, corrections, offset=0):
        """找出原始文本和校正後文本的差異

        參數:
            original_text: 原始文本
            corrected_text: 校正後的文本
            corrections: 用於存儲差異位置的列表
            offset: 文本在整體文本中的偏移量
        """
        # 如果長度相同，逐字比較
        if len(original_text) == len(corrected_text):
            i = 0
            while i < len(original_text):
                # 找出連續的不同字符
                if original_text[i] != corrected_text[i]:
                    start = i
                    while i < len(original_text) and original_text[i] != corrected_text[i]:
                        i += 1
                    corrections.append((offset + start, offset + i))
                else:
                    i += 1
        else:
            # 如果長度不同，標記整個段落
            corrections.append((offset, offset + len(corrected_text)))

    def _update_text_area(self, corrected_text, corrections=None):
        """更新文字區域的內容

        參數:
            corrected_text: 校正後的文字
            corrections: 修正的位置列表，每個元素是 (start, end) 元組
        """
        # 關閉 undo 記錄，避免校正本身被記錄
        current_undo_state = self.text_area.cget("undo")
        self.text_area.config(undo=False)

        self.text_area.delete(1.0, tk.END)
        self.text_area.insert(tk.END, corrected_text)

        # 應用紅色底線標記到修正過的文字
        if corrections:
            for start, end in corrections:
                # 將字符位置轉換為 Tkinter 的行列位置
                start_line = corrected_text[:start].count('\n') + 1
                start_col = start - corrected_text[:start].rfind('\n') - 1 if '\n' in corrected_text[:start] else start

                end_line = corrected_text[:end].count('\n') + 1
                end_col = end - corrected_text[:end].rfind('\n') - 1 if '\n' in corrected_text[:end] else end

                # 應用標籤
                try:
                    self.text_area.tag_add("corrected", f"{start_line}.{start_col}", f"{end_line}.{end_col}")
                except Exception as e:
                    print(f"應用標籤時發生錯誤: {str(e)}")

        # 恢復 undo 狀態並清空 undo 堆疊
        self.text_area.config(undo=current_undo_state)
        if current_undo_state:
            self.text_area.edit_reset()

        self.status_bar.config(text="文字校正完成")


    def clear_correction_highlights(self):
        """清除所有校正標記"""
        self.text_area.tag_remove("corrected", "1.0", tk.END)
        self.status_bar.config(text="已清除所有校正標記")

    def load_protected_words(self):
        """載入詞彙保護表

        回傳:
            詞彙保護列表
        """
        try:
            # 檢查檔案是否存在
            if not os.path.exists("protected_words.json"):
                # 如果不存在，創建一個空的詞彙保護表
                with open("protected_words.json", "w", encoding="utf-8") as f:
                    json.dump([], f, ensure_ascii=False, indent=4)
                return []

            # 讀取詞彙保護表
            with open("protected_words.json", "r", encoding="utf-8") as f:
                data = json.load(f)

            # 檢查是否為列表格式
            if isinstance(data, list):
                return data
            # 向下兼容舊格式（包含 protected_words 鍵的對象）
            elif isinstance(data, dict) and "protected_words" in data:
                return data["protected_words"]
            else:
                print("詞彙保護表格式不正確，使用空列表")
                return []

        except Exception as e:
            print(f"載入詞彙保護表時發生錯誤: {str(e)}")
            messagebox.showerror("錯誤", f"無法載入詞彙保護表: {str(e)}")
            return []

    def save_protected_words(self):
        """儲存詞彙保護表"""
        # 儲存詞彙保護表
        try:
            with open('protected_words.json', 'w', encoding='utf-8') as f:
                json.dump(self.protected_words, f, ensure_ascii=False, indent=4)
        except Exception as e:
            messagebox.showerror("錯誤", f"無法儲存詞彙保護表: {str(e)}")

    def manage_protected_words(self):
        """管理保護詞彙的視窗"""
        # 創建一個新視窗來管理保護詞彙
        manage_window = tk.Toplevel(self.root)
        manage_window.title("管理保護詞彙")
        manage_window.geometry("400x500")

        # 創建一個框架
        frame = tk.Frame(manage_window)
        frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # 添加標籤
        tk.Label(frame, text="保護詞彙列表:").pack(anchor=tk.W)

        # 添加列表框和滾動條
        list_frame = tk.Frame(frame)
        list_frame.pack(fill=tk.BOTH, expand=True)

        scrollbar = tk.Scrollbar(list_frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        words_listbox = tk.Listbox(list_frame, yscrollcommand=scrollbar.set)
        words_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.config(command=words_listbox.yview)

        # 填充列表框
        for word in self.protected_words:
            words_listbox.insert(tk.END, word)

        # 添加輸入欄位和按鈕
        input_frame = tk.Frame(frame)
        input_frame.pack(fill=tk.X, pady=5)

        tk.Label(input_frame, text="新增詞彙:").pack(side=tk.LEFT)
        word_entry = tk.Entry(input_frame)
        word_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)

        # 添加按鈕
        buttons_frame = tk.Frame(frame)
        buttons_frame.pack(fill=tk.X)

        def add_word():
            """添加新詞彙到保護列表"""
            word = word_entry.get().strip()
            if word and word not in self.protected_words:
                self.protected_words.append(word)
                words_listbox.insert(tk.END, word)
                word_entry.delete(0, tk.END)
                self.save_protected_words()

        def remove_word():
            """從保護列表中移除選中的詞彙"""
            selection = words_listbox.curselection()
            if selection:
                index = selection[0]
                word = words_listbox.get(index)
                words_listbox.delete(index)
                self.protected_words.remove(word)
                self.save_protected_words()

        def download_json():
            """下載保護詞彙為JSON檔案"""
            file_path = filedialog.asksaveasfilename(
                defaultextension=".json",
                filetypes=[("JSON檔案", "*.json"), ("所有檔案", "*.*")],
                title="儲存保護詞彙"
            )
            if file_path:
                try:
                    with open(file_path, 'w', encoding='utf-8') as f:
                        json.dump(self.protected_words, f, ensure_ascii=False, indent=4)
                    messagebox.showinfo("成功", f"已成功儲存保護詞彙至 {file_path}")
                except Exception as e:
                    messagebox.showerror("錯誤", f"儲存檔案時發生錯誤: {str(e)}")

        def upload_json():
            """從JSON檔案上傳保護詞彙"""
            file_path = filedialog.askopenfilename(
                filetypes=[("JSON檔案", "*.json"), ("所有檔案", "*.*")],
                title="選擇保護詞彙檔案"
            )
            if file_path:
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        imported_words = json.load(f)

                    # 確認匯入的資料是列表格式
                    if not isinstance(imported_words, list):
                        if isinstance(imported_words, dict) and "protected_words" in imported_words:
                            imported_words = imported_words["protected_words"]
                        else:
                            raise ValueError("檔案格式不正確，請確保是JSON陣列或包含'protected_words'鍵的物件")

                    # 計算新增的詞彙數量
                    new_words_count = 0
                    for word in imported_words:
                        if word and word not in self.protected_words:
                            self.protected_words.append(word)
                            words_listbox.insert(tk.END, word)
                            new_words_count += 1

                    # 儲存更新後的詞彙表
                    self.save_protected_words()

                    messagebox.showinfo("成功", f"已成功匯入 {new_words_count} 個新詞彙")
                except Exception as e:
                    messagebox.showerror("錯誤", f"匯入檔案時發生錯誤: {str(e)}")

        tk.Button(buttons_frame, text="添加", command=add_word).pack(side=tk.LEFT, padx=5)
        tk.Button(buttons_frame, text="刪除", command=remove_word).pack(side=tk.LEFT, padx=5)
        tk.Button(buttons_frame, text="下載JSON", command=download_json).pack(side=tk.LEFT, padx=5)
        tk.Button(buttons_frame, text="上傳JSON", command=upload_json).pack(side=tk.LEFT, padx=5)
        tk.Button(buttons_frame, text="關閉", command=manage_window.destroy).pack(side=tk.RIGHT, padx=5)

    def open_text_settings(self):
        """開啟文字格式設定視窗"""
        settings_window = tk.Toplevel(self.root)
        settings_window.title("文字格式設定")
        settings_window.geometry("400x350") # Increased height for line spacing
        settings_window.resizable(False, False)
        settings_window.transient(self.root)  # 設為主視窗的子視窗
        settings_window.grab_set()  # 模態視窗

        # 建立框架
        frame = tk.Frame(settings_window, padx=20, pady=20)
        frame.pack(fill=tk.BOTH, expand=True)

        # 字體選擇
        tk.Label(frame, text="字體:").grid(row=0, column=0, sticky=tk.W, pady=5) # Reduced pady

        # 獲取系統可用字體
        available_fonts = ["新細明體", "標楷體", "微軟正黑體", "Arial", "Times New Roman", "Courier New"]

        font_var = tk.StringVar(value=self.settings["font_family"])
        font_combo = ttk.Combobox(frame, textvariable=font_var, values=available_fonts, width=20)
        font_combo.grid(row=0, column=1, sticky=tk.W, pady=5)

        # 字體大小選擇
        tk.Label(frame, text="字體大小:").grid(row=1, column=0, sticky=tk.W, pady=5)

        size_var = tk.IntVar(value=self.settings["font_size"])
        size_combo = ttk.Combobox(frame, textvariable=size_var, values=[8, 9, 10, 11, 12, 14, 16, 18, 20, 22, 24, 26, 28, 36], width=5)
        size_combo.grid(row=1, column=1, sticky=tk.W, pady=5)

        # 行距選擇
        tk.Label(frame, text="行距 (段落後):").grid(row=2, column=0, sticky=tk.W, pady=5)
        spacing_var = tk.IntVar(value=self.settings["line_spacing"])
        # 允許 0 到 20 的間距
        spacing_spinbox = ttk.Spinbox(frame, from_=0, to=20, textvariable=spacing_var, width=5)
        spacing_spinbox.grid(row=2, column=1, sticky=tk.W, pady=5)

        # 預覽區域
        tk.Label(frame, text="預覽:").grid(row=3, column=0, sticky=tk.W, pady=5)

        preview_text = tk.Text(frame, width=30, height=4, wrap=tk.WORD)
        preview_text.grid(row=3, column=1, sticky=tk.W, pady=5)
        preview_text.insert(tk.END, "這是預覽文字\nABCDEFG\n123456789")

        # 更新預覽的函數
        def update_preview(*args):
            font_family = font_var.get()
            font_size = size_var.get()
            line_spacing = spacing_var.get()
            try:
                # 更新字體和行距
                preview_text.configure(font=(font_family, font_size), spacing3=line_spacing)
            except tk.TclError as e:
                # 處理可能的字體錯誤
                print(f"預覽錯誤: {e}")
                preview_text.configure(font=("Arial", font_size), spacing3=line_spacing) # Fallback font

        # 綁定變更事件
        font_var.trace_add("write", update_preview)
        size_var.trace_add("write", update_preview)
        spacing_var.trace_add("write", update_preview) # 綁定行距變數

        # 初始更新預覽
        update_preview()

        # 按鈕區域
        button_frame = tk.Frame(frame)
        button_frame.grid(row=4, column=0, columnspan=2, pady=15) # Adjusted row and pady

        # 確定按鈕
        def save_settings():
            self.settings["font_family"] = font_var.get()
            self.settings["font_size"] = size_var.get()
            self.settings["line_spacing"] = spacing_var.get() # 儲存行距
            self.save_settings()
            # 應用字體和行距到主文字區域
            self.text_area.configure(
                font=(self.settings["font_family"], self.settings["font_size"]),
                spacing3=self.settings["line_spacing"]
            )
            settings_window.destroy()

        tk.Button(button_frame, text="確定", command=save_settings, width=10).pack(side=tk.LEFT, padx=10)

        # 取消按鈕
        tk.Button(button_frame, text="取消", command=settings_window.destroy, width=10).pack(side=tk.LEFT, padx=10)

    def load_settings(self):
        """載入設定

        回傳:
            設定字典
        """
        default_settings = {
            "font_family": "新細明體",
            "font_size": 12,
            "line_spacing": 3, # 新增：預設行距 (段落下間距)
            "dark_mode": False
        }

        try:
            # 檢查檔案是否存在
            if not os.path.exists("settings.json"):
                # 如果不存在，創建一個預設設定檔
                with open("settings.json", "w", encoding="utf-8") as f:
                    json.dump(default_settings, f, ensure_ascii=False, indent=4)
                return default_settings

            # 讀取設定檔
            with open("settings.json", "r", encoding="utf-8") as f:
                settings = json.load(f)

            # 確保所有必要的設定都存在
            for key in default_settings:
                if key not in settings:
                    settings[key] = default_settings[key]

            return settings
        except Exception as e:
            print(f"載入設定時發生錯誤: {str(e)}")
            messagebox.showerror("錯誤", f"無法載入設定: {str(e)}")
            return default_settings

    def save_settings(self):
        """儲存設定"""
        try:
            with open("settings.json", "w", encoding="utf-8") as f:
                json.dump(self.settings, f, ensure_ascii=False, indent=4)
            print("設定已儲存")
        except Exception as e:
            print(f"儲存設定時發生錯誤: {str(e)}")
            messagebox.showerror("錯誤", f"無法儲存設定: {str(e)}")

    def toggle_dark_mode(self):
        """切換深色模式"""
        self.settings["dark_mode"] = not self.settings["dark_mode"]
        self.save_settings()
        self.apply_theme()
        self.status_bar.config(text=f"已切換至 {'深色' if self.settings['dark_mode'] else '淺色'} 模式") # 更新狀態欄提示

    def apply_theme(self):
        """應用主題設定"""
        if self.settings["dark_mode"]:
            # 深色模式
            bg_color = "#2b2b2b"
            fg_color = "white"
            text_bg = "#2b2b2b"
            text_fg = "white"
            button_bg = "#3c3f41"
            button_fg = "white"
            canvas_bg = "#2b2b2b"
            toolbar_bg = "#3c3f41" # Toolbar background
        else:
            # 淺色模式
            bg_color = "white"
            fg_color = "black"
            text_bg = "white"
            text_fg = "black"
            button_bg = "#f0f0f0"
            button_fg = "black"
            canvas_bg = "white"
            toolbar_bg = "#f0f0f0" # Toolbar background

        # 應用主題到主視窗
        self.root.configure(bg=bg_color)

        # 應用主題到文字區域
        self.text_area.configure(bg=text_bg, fg=text_fg)

        # 應用主題到圖片區域
        self.image_frame.configure(bg=bg_color)
        self.image_container.configure(bg=bg_color)
        self.image_canvas.configure(bg=canvas_bg)

        # 應用主題到圖片下載按鈕框架及其按鈕
        for widget in self.image_frame.winfo_children():
            if isinstance(widget, tk.Frame): # This is img_button_frame
                widget.configure(bg=bg_color)
                for child in widget.winfo_children():
                    if isinstance(child, tk.Button):
                        child.configure(bg=button_bg, fg=button_fg)

        # 應用主題到狀態欄
        self.status_bar.configure(bg=bg_color, fg=fg_color)

        # 應用主題到工具欄 (如果已創建)
        if hasattr(self, 'toolbar_main_frame'):
            self.toolbar_main_frame.configure(bg=toolbar_bg)
            self.toolbar_top_frame.configure(bg=toolbar_bg)
            self.toolbar_bottom_frame.configure(bg=toolbar_bg)
            # Apply to buttons in top toolbar
            for child in self.toolbar_top_frame.winfo_children():
                if isinstance(child, tk.Button):
                    child.configure(bg=button_bg, fg=button_fg)
            # Apply to widgets in bottom toolbar (if any added later)
            for child in self.toolbar_bottom_frame.winfo_children(): # 包括動態新增的按鈕
                 if isinstance(child, tk.Button):
                     child.configure(bg=button_bg, fg=button_fg)
                 elif isinstance(child, tk.Label): # Example for labels if added
                     child.configure(bg=toolbar_bg, fg=fg_color)


    def adjust_indentation(self, event=None):
        """調整文字縮進，使換行後的文字對齊前一行的第一個字"""
        # 重置修改標誌，避免無限循環
        try:
            self.text_area.edit_modified(False)
        except tk.TclError:
            # 可能在 Text 元件初始化或銷毀過程中觸發，忽略
            pass

        # 獲取所有文字
        content = self.text_area.get("1.0", tk.END)

        # 如果內容為空，不做處理
        if not content.strip():
            return

        # 處理每個段落
        lines = content.split('\n')
        for i in range(len(lines)):
            # 跳過空行
            if not lines[i].strip():
                continue

            # 獲取當前行第一個非空白字符的位置
            first_char_pos = len(lines[i]) - len(lines[i].lstrip())

            # 如果不是第一行且前一行不為空，設置縮進
            if i > 0 and lines[i-1].strip():
                prev_first_char_pos = len(lines[i-1]) - len(lines[i-1].lstrip())

                # 設置縮進標籤
                tag_name = f"indent_{i}"
                self.text_area.tag_configure(tag_name, lmargin1=prev_first_char_pos)

                # 應用標籤到當前行
                line_start = f"{i+1}.0"
                line_end = f"{i+1}.{len(lines[i])}"
                try:
                    self.text_area.tag_add(tag_name, line_start, line_end)
                except tk.TclError:
                     # 可能在 Text 元件初始化或銷毀過程中觸發，忽略
                    pass

    def adjust_text_formatting(self, event=None):
        """調整文字格式，包括縮進和對齊"""
        # 調用原有的縮進方法
        self.adjust_indentation(event)

    def handle_password_protected_file(self, file_path):
        """處理有密碼保護的Word檔案

        參數:
            file_path: 加密Word檔案的路徑
        """
        # 處理有密碼保護的檔案
        password = self.ask_password()
        if password:
            # --- 修改：調用新的處理邏輯，並傳遞密碼 ---
            self.load_and_display_word_content(file_path, password)
        else:
            # 用户取消输入密码
            self.status_bar.config(text="已取消密碼輸入")

    def ask_password(self):
        """顯示密碼輸入對話框

        回傳:
            使用者輸入的密碼
        """
        # 創建密碼輸入對話框
        password_window = tk.Toplevel(self.root)
        password_window.title("密碼保護")
        password_window.geometry("300x150")
        password_window.resizable(False, False)

        # 設置模態對話框
        password_window.transient(self.root)
        password_window.grab_set()

        # 居中顯示
        window_width = 300
        window_height = 150
        screen_width = password_window.winfo_screenwidth()
        screen_height = password_window.winfo_screenheight()
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2
        password_window.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 添加說明標籤
        tk.Label(password_window, text="該檔案有密碼保護，請輸入密碼:", font=("Arial", 10)).pack(pady=10)

        # 密碼輸入框
        password_entry = tk.Entry(password_window, show="*", width=25)
        password_entry.pack(pady=5)
        password_entry.focus_set()  # 設置焦點

        password = None

        # 確定按鈕回調函數
        def on_ok():
            nonlocal password
            password = password_entry.get()
            password_window.destroy()

        # 取消按鈕回調函數
        def on_cancel():
            password_window.destroy()

        # 按鈕區域
        button_frame = tk.Frame(password_window)
        button_frame.pack(pady=10)

        tk.Button(button_frame, text="確定", command=on_ok, width=10).pack(side=tk.LEFT, padx=5)
        tk.Button(button_frame, text="取消", command=on_cancel, width=10).pack(side=tk.LEFT, padx=5)

        # 綁定回車鍵
        password_window.bind("<Return>", lambda event: on_ok())
        password_window.bind("<Escape>", lambda event: on_cancel())

        # 等待視窗關閉
        password_window.wait_window()
        return password

    def log_error(self, error_type, error_message, details=None):
        """記錄錯誤到日誌檔案

        參數:
            error_type: 錯誤類型
            error_message: 錯誤訊息
            details: 詳細錯誤信息（可選）
        """
        try:
            error_log = f"錯誤類型: {error_type}\n錯誤訊息: {error_message}"
            if details:
                error_log += f"\n詳細信息: {details}"

            # 記錄到日誌檔案
            logging.error(error_log)

            # 顯示錯誤訊息給使用者
            messagebox.showerror("錯誤", f"{error_message}\n\n錯誤已記錄到日誌檔案中。")
        except Exception as e:
            # 如果記錄錯誤時發生錯誤，直接顯示訊息
            messagebox.showerror("錯誤", f"無法記錄錯誤: {str(e)}\n原始錯誤: {error_message}")

    def view_error_logs(self):
        """檢視錯誤日誌"""
        # 創建錯誤日誌視窗
        log_window = tk.Toplevel(self.root)
        log_window.title("錯誤日誌檢視")
        log_window.geometry("800x500")
        log_window.transient(self.root)  # 設為主視窗的子視窗
        log_window.grab_set()  # 模態視窗

        # 創建框架
        frame = tk.Frame(log_window, padx=10, pady=10)
        frame.pack(fill=tk.BOTH, expand=True)

        # 日誌檔案列表
        tk.Label(frame, text="選擇日誌檔案:").pack(anchor=tk.W, pady=(0, 5))

        # 獲取日誌檔案列表
        log_dir = "logs"
        if not os.path.exists(log_dir):
            os.makedirs(log_dir)

        log_files = [f for f in os.listdir(log_dir) if f.startswith("error_log_") and f.endswith(".log")]
        log_files.sort(reverse=True)  # 最新的日誌檔案排在前面

        if not log_files:
            tk.Label(frame, text="沒有找到錯誤日誌檔案。").pack(pady=20)
            tk.Button(frame, text="關閉", command=log_window.destroy).pack(pady=10)
            return

        # 日誌檔案下拉選單
        selected_log = tk.StringVar(value=log_files[0] if log_files else "")
        log_combo = ttk.Combobox(frame, textvariable=selected_log, values=log_files, width=40, state="readonly")
        log_combo.pack(anchor=tk.W, pady=(0, 10))

        # 日誌內容顯示區域
        tk.Label(frame, text="日誌內容:").pack(anchor=tk.W, pady=(0, 5))

        # 添加滾動條
        scrollbar = tk.Scrollbar(frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        # 日誌內容文字區域
        log_text = tk.Text(frame, wrap=tk.WORD, yscrollcommand=scrollbar.set)
        log_text.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        scrollbar.config(command=log_text.yview)

        # 更新日誌內容的函數
        def update_log_content(*args):
            log_text.delete(1.0, tk.END)  # 清空文字區域
            selected_file = selected_log.get()

            if not selected_file:
                return

            try:
                with open(os.path.join(log_dir, selected_file), "r", encoding="utf-8") as f:
                    content = f.read()
                    if content:
                        log_text.insert(tk.END, content)
                    else:
                        log_text.insert(tk.END, "日誌檔案為空。")
            except Exception as e:
                log_text.insert(tk.END, f"無法讀取日誌檔案: {str(e)}")

        # 綁定選擇事件
        log_combo.bind("<<ComboboxSelected>>", update_log_content)

        # 按鈕區域
        button_frame = tk.Frame(frame)
        button_frame.pack(fill=tk.X, pady=(10, 0))

        # 刪除日誌按鈕
        def delete_log():
            selected_file = selected_log.get()
            if not selected_file:
                return

            if messagebox.askyesno("確認刪除", f"確定要刪除日誌檔案 {selected_file} 嗎？"):
                try:
                    os.remove(os.path.join(log_dir, selected_file))
                    # 更新日誌檔案列表
                    log_files = [f for f in os.listdir(log_dir) if f.startswith("error_log_") and f.endswith(".log")]
                    log_files.sort(reverse=True)
                    log_combo.config(values=log_files)

                    if log_files:
                        selected_log.set(log_files[0])
                        update_log_content()
                    else:
                        selected_log.set("")
                        log_text.delete(1.0, tk.END)
                        log_text.insert(tk.END, "沒有找到錯誤日誌檔案。")
                except Exception as e:
                    messagebox.showerror("錯誤", f"無法刪除日誌檔案: {str(e)}")

        tk.Button(button_frame, text="刪除日誌", command=delete_log).pack(side=tk.LEFT, padx=(0, 10))
        tk.Button(button_frame, text="關閉", command=log_window.destroy).pack(side=tk.RIGHT)

        # 初始顯示第一個日誌檔案的內容
        update_log_content()

    # --- 新增：從 test_01.py 複製來的 COM 解析函數 (修正縮排) ---
    def parse_word_document_com(self, filepath: str):
        """
        使用 Windows COM 與 Microsoft Word 互動來解析 .docx 文件，
        以嘗試獲取包括自動編號在內的渲染後文字。

        Args:
            filepath (str): Word 文件的路徑。

        Returns:
            str | None: 解析後的文字內容，包含自動編號和縮排。
                          如果發生錯誤、缺少依賴或無法執行則返回 None。
        """
        # 確保路徑是 Path 對象
        filepath = Path(filepath)

        if not platform.system() == 'Windows':
            print("錯誤：COM 功能僅在 Windows 上受支持。")
            return None
        if not HAS_PYWIN32:
            print("錯誤：缺少 pywin32 模組，無法使用 COM 功能。")
            return None

        if not filepath.is_file():
            print(f"錯誤：找不到檔案 {filepath}")
            return None

        word_app = None
        doc = None
        com_initialized = False
        try:
            # 初始化 COM 環境
            try:
                pythoncom.CoInitialize()
                com_initialized = True
            except Exception as e:
                print(f"警告：初始化 COM 失敗: {e} (嘗試繼續)")

            parsed_content = []
            # 啟動 Word 應用程式 (嘗試後台執行)
            try:
                word_app = win32.Dispatch("Word.Application")
            except pythoncom.com_error as ce:
                 hr, msg, exc, arg = ce.args
                 print(f"COM 錯誤 (Dispatch): HRESULT={hr}, Message={msg}")
                 if hr == -2147221005: # CO_E_CLASSSTRING
                     messagebox.showerror("COM 錯誤", "無法啟動 Microsoft Word。\n請確認已正確安裝 Word。")
                 else:
                     messagebox.showerror("COM 錯誤", f"啟動 Word 時發生 COM 錯誤:\n{msg}")
                 return None

            word_app.Visible = False # 不顯示 Word 視窗
            try:
                word_app.DisplayAlerts = 0 # wdAlertsNone = 0
            except Exception as e:
                print(f"警告：無法設置 DisplayAlerts 屬性: {e}")

            # 打開文件
            try:
                 doc = word_app.Documents.Open(str(filepath.resolve()), ReadOnly=True)
            except pythoncom.com_error as ce:
                hr, msg, exc, arg = ce.args
                print(f"COM 錯誤 (Open): HRESULT={hr}, Message={msg}")
                messagebox.showerror("文件開啟錯誤", f"無法透過 Word 開啟檔案 '{filepath.name}':\n{msg}\n\n請檢查文件是否存在、未損壞且 Word 可以開啟它。")
                if word_app:
                    try: word_app.Quit()
                    except: pass
                return None

            # --- 處理縮排 ---
            POINTS_PER_INDENT_LEVEL = 18 # 每 18 磅 (Point) 算一級縮排 (可調整)
            SPACES_PER_INDENT_LEVEL = 3  # 每級縮排對應的空格數 (可調整)

            # 迭代文件中的段落
            try:
                for i, para_com in enumerate(doc.Paragraphs):
                    indent_space = ""
                    formatted_line = "[讀取段落時發生錯誤]"
                    try:
                        para_range = para_com.Range
                        list_string = para_range.ListFormat.ListString
                        full_text = para_range.Text
                        actual_text = full_text.rstrip('\r\n')

                        # --- 計算縮排 ---
                        indent_points = 0.0
                        try:
                            indent_points = para_com.Format.LeftIndent
                        except AttributeError: pass
                        except pythoncom.com_error as ce: print(f"警告：獲取段落 {i+1} 縮排時 COM 錯誤: {ce}")
                        except Exception as indent_err: print(f"警告：獲取段落 {i+1} 縮排時出錯: {indent_err}")

                        if indent_points > 0:
                            indent_level = int(indent_points / POINTS_PER_INDENT_LEVEL)
                            if indent_level < 0: indent_level = 0
                            indent_space = " " * (indent_level * SPACES_PER_INDENT_LEVEL)

                        # --- 組合輸出 ---
                        separator = "\t"

                        if list_string:
                            temp_text = actual_text
                            if list_string and temp_text.startswith(list_string):
                                temp_text = temp_text[len(list_string):]
                                temp_text = temp_text.lstrip(' \t')
                            formatted_line = f"{indent_space}{list_string}{separator}{temp_text}"
                        else:
                            formatted_line = f"{indent_space}{actual_text}"

                        parsed_content.append(formatted_line)

                    except pythoncom.com_error as para_ce:
                         print(f"警告：讀取段落 {i+1} 時發生 COM 錯誤: {para_ce}")
                         parsed_content.append(f"{indent_space}[讀取段落 COM 錯誤]")
                    except Exception as para_exc:
                        print(f"警告：讀取段落 {i+1} 時發生未知錯誤: {para_exc}")
                        try:
                            raw_text = para_com.Range.Text.rstrip('\r\n')
                            parsed_content.append(f"{indent_space}[讀取錯誤] {raw_text}")
                        except:
                            parsed_content.append(f"{indent_space}[讀取錯誤且無法獲取原始文本]")

            except Exception as iter_exc:
                print(f"錯誤：迭代段落時發生嚴重錯誤: {iter_exc}\n{traceback.format_exc()}")
                messagebox.showerror("解析錯誤", f"處理文件 '{filepath.name}' 段落時發生錯誤:\n{iter_exc}\n\n可能只能顯示部分內容。")
                pass

            # 添加提示：不顯示圖片 (圖片提取由 extract_images_from_docx 處理)
            # parsed_content.append("\n\n--- (注意：文件中的圖片需另外提取) ---")

            return "\n".join(parsed_content)

        except pythoncom.com_error as ce:
            hr, msg, exc, arg = ce.args
            print(f"處理 Word 文件時發生 COM 錯誤: HRESULT={hr}, Message={msg}\n{traceback.format_exc()}")
            messagebox.showerror("COM 交互錯誤", f"與 Word 交互時發生 COM 錯誤：\n{msg}\n(請確認 Word 可正常運作)")
            return None
        except Exception as e:
            error_details = traceback.format_exc()
            print(f"解析 Word 文件時發生未知錯誤: {e}\n{error_details}")
            messagebox.showerror("未知解析錯誤", f"解析 Word 文件時發生未知錯誤：\n{e}\n(詳細資訊請查看控制台輸出)")
            return None
        finally:
            # --- 清理 COM ---
            try:
                if doc:
                    doc.Close(SaveChanges=0)
                    doc = None
            except Exception as e_close: print(f"關閉 Word 文件時發生錯誤: {e_close}")
            try:
                if word_app:
                    word_app.Quit()
                    word_app = None
            except Exception as e_quit: print(f"退出 Word 應用程式時發生錯誤: {e_quit}")
            if com_initialized:
                try:
                    pythoncom.CoUninitialize()
                except Exception as e_uninit: print(f"警告：清理 COM 環境失敗: {e_uninit}")

# End of TextCorrectionTool class definition


# --- Main execution block at top level ---
def main():
    """程式主入口點"""
    try:
        # 嘗試使用 TkinterDnD2 創建支援拖放的根視窗
        try:
            # 延遲導入 tkinterdnd2，僅在需要時導入
            from tkinterdnd2 import TkinterDnD
            root = TkinterDnD.Tk()
            print("成功使用 TkinterDnD2 初始化根視窗")
        except Exception as e:
            print(f"無法使用 TkinterDnD2: {str(e)}")
            # 退回使用普通的 Tk
            root = tk.Tk()
            print("使用普通 Tk 初始化根視窗")

        app = TextCorrectionTool(root)
        root.mainloop()
    except Exception as e:
        error_msg = f"程式執行時發生嚴重錯誤: {str(e)}\n{traceback.format_exc()}"
        print(error_msg)
        # 嘗試記錄錯誤
        try:
            logging.error(error_msg)
        except Exception as log_e:
            print(f"記錄錯誤時也發生錯誤: {log_e}")
        # 嘗試顯示錯誤訊息框
        try:
            messagebox.showerror("嚴重錯誤", f"程式執行時發生嚴重錯誤: {str(e)}\n請查看日誌檔案獲取詳細信息。")
        except Exception as msg_e:
             print(f"顯示錯誤訊息框時也發生錯誤: {msg_e}")

if __name__ == "__main__":
    main()
